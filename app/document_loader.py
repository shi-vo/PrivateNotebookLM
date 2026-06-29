"""
Extract raw text from uploaded documents.

Every loader returns a list of "units" with the shape:

    {"text": str, "page_or_section": str, "doc_id": str}

- PDFs: one unit per page, page_or_section = "Page <n>" (1-indexed).
- DOCX: paragraphs are grouped into sections (a section break = a
  paragraph that looks like a heading, or every N paragraphs as a
  fallback) so we don't end up with one "unit" per single sentence.
  page_or_section = "Section <n>".
- TXT: split on blank lines into pseudo-paragraphs, grouped the same
  way as DOCX. page_or_section = "Section <n>".

The chunker (chunker.py) later splits/merges these units into
token-bounded chunks, carrying page_or_section along as metadata.
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import List, TypedDict

logger = logging.getLogger(__name__)


class TextUnit(TypedDict):
    text: str
    page_or_section: str
    doc_id: str


class UnsupportedFileType(Exception):
    pass


class EmptyDocumentError(Exception):
    """Raised when extraction yields no usable text at all."""


def load_document(file_path: Path, doc_id: str) -> List[TextUnit]:
    """Dispatch to the right loader based on file extension."""
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        units = _load_pdf(file_path, doc_id)
    elif ext == ".docx":
        units = _load_docx(file_path, doc_id)
    elif ext == ".txt":
        units = _load_txt(file_path, doc_id)
    else:
        raise UnsupportedFileType(
            f"Unsupported file type '{ext}'. Supported: .pdf, .docx, .txt"
        )

    units = [u for u in units if u["text"].strip()]
    if not units:
        raise EmptyDocumentError(
            f"No extractable text found in '{file_path.name}'. "
            "The file may be empty, image-only (scanned), or corrupted."
        )
    return units


# --------------------------------------------------------------------------
# PDF
# --------------------------------------------------------------------------
def _load_pdf(file_path: Path, doc_id: str) -> List[TextUnit]:
    """Page-by-page extraction with pypdf, falling back to pdfplumber for
    pages pypdf can't get usable text from (tricky layouts, columns,
    tables)."""
    units: List[TextUnit] = []
    pypdf_failed_pages: List[int] = []

    try:
        from pypdf import PdfReader

        reader = PdfReader(str(file_path))
        for i, page in enumerate(reader.pages):
            page_num = i + 1
            try:
                text = page.extract_text() or ""
            except Exception as e:  # pragma: no cover - defensive
                logger.warning("pypdf failed on page %d: %s", page_num, e)
                text = ""

            if text.strip():
                units.append(
                    {
                        "text": text.strip(),
                        "page_or_section": f"Page {page_num}",
                        "doc_id": doc_id,
                    }
                )
            else:
                pypdf_failed_pages.append(page_num)
    except Exception as e:
        logger.warning("pypdf could not open '%s': %s. Falling back to pdfplumber for the whole file.", file_path.name, e)
        pypdf_failed_pages = None  # signal: redo every page with pdfplumber

    # Fallback: any page pypdf couldn't extract (or the whole file, if
    # pypdf couldn't even open it) gets a second attempt with pdfplumber,
    # which handles some layouts/tables better.
    if pypdf_failed_pages:
        units.extend(_load_pdf_pages_with_pdfplumber(file_path, doc_id, pypdf_failed_pages))
    elif pypdf_failed_pages is None:
        units = _load_pdf_pages_with_pdfplumber(file_path, doc_id, None)

    # Re-sort by page number since fallback pages may have been appended
    # out of order.
    def _page_num(u: TextUnit) -> int:
        try:
            return int(u["page_or_section"].split()[-1])
        except (ValueError, IndexError):
            return 0

    units.sort(key=_page_num)
    return units


def _load_pdf_pages_with_pdfplumber(
    file_path: Path, doc_id: str, page_numbers: List[int] | None
) -> List[TextUnit]:
    """page_numbers is 1-indexed list of pages to (re)extract, or None for all."""
    units: List[TextUnit] = []
    try:
        import pdfplumber

        with pdfplumber.open(str(file_path)) as pdf:
            targets = (
                range(len(pdf.pages))
                if page_numbers is None
                else [p - 1 for p in page_numbers]
            )
            for i in targets:
                if i < 0 or i >= len(pdf.pages):
                    continue
                page_num = i + 1
                try:
                    text = pdf.pages[i].extract_text() or ""
                except Exception as e:  # pragma: no cover - defensive
                    logger.warning("pdfplumber failed on page %d: %s", page_num, e)
                    text = ""
                if text.strip():
                    units.append(
                        {
                            "text": text.strip(),
                            "page_or_section": f"Page {page_num}",
                            "doc_id": doc_id,
                        }
                    )
    except Exception as e:
        logger.error("pdfplumber fallback also failed for '%s': %s", file_path.name, e)
    return units


# --------------------------------------------------------------------------
# DOCX
# --------------------------------------------------------------------------
def _load_docx(file_path: Path, doc_id: str) -> List[TextUnit]:
    """Group paragraphs into sections. A new section starts whenever a
    paragraph uses a Heading style, or after every 8 plain paragraphs as
    a fallback so very long documents with no headings still get split
    into multiple sections."""
    from docx import Document

    doc = Document(str(file_path))

    sections: List[List[str]] = []
    current: List[str] = []
    plain_para_count = 0
    MAX_PLAIN_PARAS_PER_SECTION = 8

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = (para.style.name if para.style else "") or ""
        is_heading = style_name.lower().startswith("heading") or style_name.lower() == "title"

        if is_heading and current:
            sections.append(current)
            current = [text]
            plain_para_count = 0
            continue

        current.append(text)
        if not is_heading:
            plain_para_count += 1
            if plain_para_count >= MAX_PLAIN_PARAS_PER_SECTION:
                sections.append(current)
                current = []
                plain_para_count = 0

    if current:
        sections.append(current)

    # Also pull text out of tables (contracts often have key terms in
    # tables) as additional sections appended at the end.
    for t_idx, table in enumerate(doc.tables):
        rows_text = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                rows_text.append(" | ".join(cells))
        if rows_text:
            sections.append([f"[Table {t_idx + 1}]"] + rows_text)

    units: List[TextUnit] = []
    for i, section_paras in enumerate(sections):
        text = "\n".join(section_paras).strip()
        if text:
            units.append(
                {
                    "text": text,
                    "page_or_section": f"Section {i + 1}",
                    "doc_id": doc_id,
                }
            )
    return units


# --------------------------------------------------------------------------
# TXT
# --------------------------------------------------------------------------
def _load_txt(file_path: Path, doc_id: str) -> List[TextUnit]:
    """Split on blank lines into paragraphs, grouped into sections of up
    to 8 paragraphs each (same heuristic as DOCX, since plain text has no
    heading information at all)."""
    raw = file_path.read_text(encoding="utf-8", errors="replace")
    paragraphs = [p.strip() for p in raw.split("\n\n") if p.strip()]

    if not paragraphs and raw.strip():
        # No blank-line breaks at all — treat the whole file as one paragraph.
        paragraphs = [raw.strip()]

    MAX_PARAS_PER_SECTION = 8
    units: List[TextUnit] = []
    for i in range(0, len(paragraphs), MAX_PARAS_PER_SECTION):
        group = paragraphs[i : i + MAX_PARAS_PER_SECTION]
        section_num = (i // MAX_PARAS_PER_SECTION) + 1
        units.append(
            {
                "text": "\n\n".join(group),
                "page_or_section": f"Section {section_num}",
                "doc_id": doc_id,
            }
        )
    return units
